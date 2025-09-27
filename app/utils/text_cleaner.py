import io
import os
import re
import requests
import openpyxl # type: ignore
import pytesseract # type: ignore
import pdfplumber # type: ignore
import cv2 # type: ignore # Dependencia: opencv-python
import numpy as np # type: ignore
from PIL import Image # type: ignore
from docx import Document # type: ignore
from pdf2image import convert_from_bytes # type: ignore
from defusedxml import ElementTree as ET # type: ignore # Para parseo seguro de XML

# Límite de caracteres para el contenido final. Ajusta según tus necesidades.
MAX_FILE_TEXT_LENGTH = 4000

class AdvancedFileCleaner:
    """
    Clase para descargar y extraer texto limpio de varios tipos de archivos.
    """

    def _preprocess_image_for_ocr(self, image_bytes: bytes) -> Image.Image:
        """
        Preprocesa una imagen para mejorar la precisión del OCR.
        Convierte a escala de grises, aplica umbral adaptativo y reduce el ruido.
        """
        try:
            # Lee la imagen desde los bytes
            image_np = np.frombuffer(image_bytes, np.uint8)
            image = cv2.imdecode(image_np, cv2.IMREAD_COLOR)

            # 1. Convertir a escala de grises
            gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) # type: ignore

            # 2. Aplicar umbral adaptativo para binarizar la imagen
            # Esto ayuda a manejar diferentes condiciones de iluminación.
            processed_image = cv2.adaptiveThreshold(
                gray_image, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY, 11, 2
            )

            # Opcional: aplicar una pequeña reducción de ruido
            processed_image = cv2.medianBlur(processed_image, 3)

            # Convertir de vuelta a un objeto PIL.Image para Pytesseract
            return Image.fromarray(processed_image)
        except Exception:
            # Si OpenCV falla, regresa la imagen original desde bytes
            return Image.open(io.BytesIO(image_bytes))

    def _perform_ocr(self, image_bytes: bytes) -> str:
        """
        Realiza OCR en una imagen después de preprocesarla.
        """
        try:
            preprocessed_image = self._preprocess_image_for_ocr(image_bytes)
            return pytesseract.image_to_string(preprocessed_image) # Especificar idioma mejora la precisión
        except pytesseract.TesseractNotFoundError:
            print("ERROR: Tesseract no está instalado o no se encuentra en el PATH.")
            return ""
        except Exception as e:
            print(f"Error durante el OCR: {e}")
            return ""

    def _post_process_text(self, text: str) -> str:
        """
        Limpia el texto extraído eliminando sellos digitales, cadenas largas,
        y normalizando los espacios en blanco.
        """
        if not text:
            return ""

        # 1. Patrones de Regex para eliminar ruido (sellos, cadenas de certificación, etc.)
        # re.IGNORECASE para no distinguir mayúsculas/minúsculas, re.DOTALL para que '.' incluya saltos de línea.
        patterns_to_remove = [
            # Patrones específicos para CFDI/SAT
            r"Sello\s*digital\s*del\s*CFDI\s*:\s*[a-zA-Z0-9+/=\s]+",
            r"Sello\s*digital\s*del\s*SAT\s*:\s*[a-zA-Z0-9+/=\s]+",
            r"Cadena\s*Original\s*del\s*complemento\s*de\s*certificación\s*digital\s*del\s*SAT\s*:\s*\|\|1\.[0-9]\|[a-zA-Z0-9\-|\s\.]+",

            # Patrón más general para cadenas largas alfanuméricas (tipo Base64, hashes)
            # Busca palabras de 40+ caracteres que son casi exclusivamente alfanuméricos y algunos símbolos.
            r'\b[a-zA-Z0-9+/=\-]{40,}\b',

            # Patrón para UUIDs
            r'\b[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}\b',
        ]

        for pattern in patterns_to_remove:
            text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.DOTALL)

        # 2. Normalizar espacios en blanco y saltos de línea
        # Reemplazar múltiples espacios/tabs/saltos de línea con un solo espacio
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Opcional: si prefieres mantener saltos de línea entre párrafos, usa una lógica más compleja:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        text = "\n".join(lines)
        text = re.sub(r'\n{2,}', '\n\n', text) # Limitar a un máximo de un salto de línea doble

        return text

    def _extract_from_pdf(self, content: bytes) -> str:
        """Extrae texto de un archivo PDF, con fallback a OCR mejorado."""
        text_parts = []
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)

                # Si no se extrajo texto (PDF escaneado), intentar OCR página por página
                if not text_parts:
                    for page in pdf.pages:
                        # Aumentar resolución para mejor calidad de imagen
                        im = page.to_image(resolution=300)
                        ocr_text = self._perform_ocr(im.original_bytes) # type: ignore
                        if ocr_text and ocr_text.strip():
                            text_parts.append(ocr_text)
        except Exception as e:
            print(f"Error procesando PDF con pdfplumber: {e}. Intentando OCR directo.")
            # Fallback si pdfplumber falla: OCR directo sobre el contenido
            try:
                images = convert_from_bytes(content)
                for image in images:
                    img_byte_arr = io.BytesIO()
                    image.save(img_byte_arr, format='PNG')
                    ocr_text = self._perform_ocr(img_byte_arr.getvalue())
                    if ocr_text and ocr_text.strip():
                        text_parts.append(ocr_text)
            except Exception as ocr_e:
                print(f"Error en fallback de OCR para PDF: {ocr_e}")
        
        return "\n".join(text_parts)

    def _extract_from_xml(self, content: bytes) -> str:
        """Extrae texto de un archivo XML de forma segura."""
        try:
            # Usar defusedxml para prevenir vulnerabilidades
            root = ET.fromstring(content)
            # Extraer todo el texto de manera más limpia
            return ' '.join(node.strip() for node in root.itertext() if node.strip())
        except ET.ParseError as e:
            print(f"Error de parseo en XML: {e}")
            return ""

    def _extract_from_image(self, content: bytes) -> str:
        """Extrae texto de una imagen usando OCR mejorado."""
        return self._perform_ocr(content)

    def _extract_from_xlsx(self, content: bytes) -> str:
        """Extrae texto de un archivo XLSX, incluyendo todas las hojas."""
        text_parts = []
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(content))
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"--- Hoja: {sheet_name} ---")
                for row in sheet.iter_rows():
                    row_values = [str(cell.value) for cell in row if cell.value is not None]
                    if row_values:
                        text_parts.append("\t".join(row_values))
        except Exception as e:
            print(f"Error procesando XLSX: {e}")
        return "\n".join(text_parts)

    def _extract_from_docx(self, content: bytes) -> str:
        """Extrae texto de un archivo DOCX, incluyendo párrafos, tablas, encabezados y pies de página."""
        text_parts = []
        try:
            document = Document(io.BytesIO(content))
            
            # Extraer texto de encabezados y pies de página
            for section in document.sections:
                for header in section.header.paragraphs:
                    if header.text.strip(): text_parts.append(header.text)
                for footer in section.footer.paragraphs:
                    if footer.text.strip(): text_parts.append(footer.text)

            # Extraer texto de párrafos y tablas en el cuerpo
            for block in document.element.body:
                if block.tag.endswith('p'): # Es un párrafo
                    para = next((p for p in document.paragraphs if p._p is block), None)
                    if para and para.text.strip():
                        text_parts.append(para.text)
                elif block.tag.endswith('tbl'): # Es una tabla
                    table = next((t for t in document.tables if t._tbl is block), None)
                    if table:
                        for row in table.rows:
                            row_text = "\t".join([cell.text.strip() for cell in row.cells])
                            if row_text:
                                text_parts.append(row_text)
        except Exception as e:
            print(f"Error procesando DOCX: {e}")
        return "\n".join(text_parts)

    def _extract_from_txt(self, content: bytes) -> str:
        """Extrae texto de un archivo TXT, probando varias codificaciones."""
        encodings_to_try = ['utf-8', 'latin-1', 'iso-8859-1', 'windows-1252']
        for encoding in encodings_to_try:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        print("No se pudo decodificar el archivo TXT con las codificaciones comunes.")
        return ""

    def get_file_content_as_text(self, public_url: str, file_name: str) -> str:
        """
        Descarga un archivo y extrae su contenido como texto limpio y procesado.
        """
        if not public_url:
            return ""

        try:
            response = requests.get(public_url, stream=True, timeout=30)
            response.raise_for_status()
            content = response.content
        except requests.exceptions.RequestException as e:
            print(f"Error al descargar el archivo {public_url}: {e}")
            return ""

        file_extension = os.path.splitext(file_name.lower())[1]
        raw_text = ""

        # Mapeo de extensiones a funciones de extracción
        extraction_map = {
            '.pdf': self._extract_from_pdf,
            '.xml': self._extract_from_xml,
            '.png': self._extract_from_image,
            '.jpg': self._extract_from_image,
            '.jpeg': self._extract_from_image,
            '.gif': self._extract_from_image,
            '.bmp': self._extract_from_image,
            '.xlsx': self._extract_from_xlsx,
            '.docx': self._extract_from_docx,
            '.txt': self._extract_from_txt,
        }

        handler = extraction_map.get(file_extension)
        if handler:
            raw_text = handler(content)
        else:
            print(f"Extensión de archivo no soportada: {file_extension}")
            return ""

        # Limpieza final y truncamiento
        cleaned_text = self._post_process_text(raw_text)
        return cleaned_text[:MAX_FILE_TEXT_LENGTH]
